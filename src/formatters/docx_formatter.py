from typing import List, Any
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DocxFormatter:
    """Класс для форматирования документов в формате docx"""

    def __init__(self):
        self.document = None

    def save_exercises(self, exercises: List[Any], filename: str) -> None:
        """Сохраняет упражнения в docx файл"""
        self.document = Document()

        # Заголовок
        title = self.document.add_heading('Сборник упражнений', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Группируем упражнения по типу
        exercises_by_type = {}
        for ex in exercises:
            ex_type = ex.__class__.__name__
            if ex_type not in exercises_by_type:
                exercises_by_type[ex_type] = []
            exercises_by_type[ex_type].append(ex)

        # Добавляем упражнения по разделам
        for ex_type, ex_list in exercises_by_type.items():
            # Название раздела
            self.document.add_heading(self._get_type_name(ex_type), level=1)

            for i, ex in enumerate(ex_list, 1):
                # Номер упражнения
                self.document.add_heading(f"Упражнение {i}", level=2)

                # Описание
                self.document.add_paragraph(f"Задание: {ex.description}")

                # Вопрос
                p = self.document.add_paragraph()
                p.add_run("Вопрос: ").bold = True
                p.add_run(ex.question)

                # Специальное форматирование для разных типов упражнений
                if ex_type == 'MatchingExercise':
                    self._format_matching_exercise(ex)
                elif ex_type == 'TrueFalseExercise':
                    self._format_truefalse_exercise(ex)
                elif ex.options:  # Для остальных типов с вариантами ответов
                    p = self.document.add_paragraph()
                    p.add_run("Варианты ответов:").bold = True
                    for opt in ex.options:
                        self.document.add_paragraph(f"• {opt}", style='List Bullet')

                # Разделитель
                self.document.add_paragraph("─" * 50)

        # Сохраняем
        self.document.save(filename)

    def save_answers(self, exercises: List[Any], filename: str) -> None:
        """Сохраняет ответы в docx файл"""
        self.document = Document()

        # Заголовок
        title = self.document.add_heading('Ключи к упражнениям', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for i, ex in enumerate(exercises, 1):
            # Ответ на упражнение
            p = self.document.add_paragraph()
            p.add_run(f"Упражнение {i} ({self._get_type_name(ex.__class__.__name__)}): ").bold = True

            # Форматируем ответ в зависимости от типа упражнения
            if ex.__class__.__name__ == 'MatchingExercise':
                self._format_matching_answer_inline(ex, p)
            elif ex.__class__.__name__ == 'TrueFalseExercise':
                self._format_truefalse_answer(ex, p)
            else:
                p.add_run(str(ex.answer))

            # Добавляем пояснения для TrueFalseExercise
            if ex.__class__.__name__ == 'TrueFalseExercise' and hasattr(ex, 'get_explanations'):
                explanations = ex.get_explanations()
                if explanations:
                    p = self.document.add_paragraph()
                    p.add_run("Пояснения:").bold = True
                    for j, explanation in enumerate(explanations, 1):
                        self.document.add_paragraph(f"{j}. {explanation}", style='List Bullet')

            # Разделитель между упражнениями
            self.document.add_paragraph()

        self.document.save(filename)

    def _get_type_name(self, class_name: str) -> str:
        """Преобразует имя класса в читаемое название"""
        names = {
            'WordOrderExercise': 'Порядок слов',
            'FillBlanksExercise': 'Заполнение пропусков',
            'MultipleChoiceExercise': 'Множественный выбор',
            'MatchingExercise': 'Соответствие',
            'TrueFalseExercise': 'Верно/Неверно'
        }
        return names.get(class_name, class_name)

    def _format_matching_exercise(self, exercise) -> None:
        """Форматирует упражнение на соответствие с таблицей"""
        # Создаем таблицу для левого столбца
        if hasattr(exercise, 'left_column') and exercise.left_column:
            table = self.document.add_table(rows=len(exercise.left_column) + 1, cols=2)
            table.style = 'Table Grid'

            # Заголовки
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "№"
            hdr_cells[1].text = "Слова"

            # Заполняем левый столбец
            for idx, word in enumerate(exercise.left_column, 1):
                row_cells = table.rows[idx].cells
                row_cells[0].text = str(idx)
                row_cells[1].text = word

        # Добавляем правый столбец с определениями
        if hasattr(exercise, 'right_column') and exercise.right_column:
            self.document.add_paragraph()
            p = self.document.add_paragraph()
            p.add_run("Найдите соответствия среди определений:").bold = True

            for idx, definition in enumerate(exercise.right_column, 1):
                self.document.add_paragraph(f"{chr(64 + idx)}. {definition}", style='List Bullet')

        # Добавляем инструкцию
        self.document.add_paragraph()
        self.document.add_paragraph("Запишите соответствия в формате: 1-А, 2-В и т.д.")

    def _format_truefalse_exercise(self, exercise) -> None:
        """Форматирует упражнение Верно/Неверно с таблицей"""
        if hasattr(exercise, 'statements') and exercise.statements:
            # Создаем таблицу для утверждений
            table = self.document.add_table(rows=len(exercise.statements) + 1, cols=3)
            table.style = 'Table Grid'

            # Заголовки
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "№"
            hdr_cells[1].text = "Утверждение"
            hdr_cells[2].text = "Ответ (В/Н)"

            # Заполняем утверждения
            for idx, stmt in enumerate(exercise.statements, 1):
                row_cells = table.rows[idx].cells
                row_cells[0].text = str(idx)
                row_cells[1].text = stmt['text']
                row_cells[2].text = "_____"

    def _format_matching_answer_inline(self, exercise, paragraph) -> None:
        """Форматирует ответ для matching упражнения в одну строку"""
        if hasattr(exercise, 'pairs'):
            paragraph.add_run("\n")
            for k, v in exercise.pairs.items():
                paragraph.add_run(f"{k} → {v}\n")

    def _format_truefalse_answer(self, exercise, paragraph) -> None:
        """Форматирует ответ для true/false упражнения"""
        if hasattr(exercise, 'statements') and exercise.statements:
            paragraph.add_run("\n")
            for idx, stmt in enumerate(exercise.statements, 1):
                answer = "Верно" if stmt['is_true'] else "Неверно"
                paragraph.add_run(f"{idx}. {answer}\n")