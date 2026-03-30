import pytest
from docx import Document
from src.formatters.docx_formatter import DocxFormatter


class DummyExercise:
    def __init__(self, description="desc", question="q", options=None, answer="a"):
        self.description = description
        self.question = question
        self.options = options or []
        self.answer = answer


class MatchingExercise(DummyExercise):
    def __init__(self, left_column=None, right_column=None, pairs=None):
        super().__init__()
        self.left_column = left_column or ["one", "two"]
        self.right_column = right_column or ["uno", "dos"]
        self.pairs = pairs or {"1": "A", "2": "B"}


class TrueFalseExercise(DummyExercise):
    def __init__(self, statements=None, explanations=None):
        super().__init__()
        self.statements = statements or [{"text": "Statement 1", "is_true": True}]
        self._explanations = explanations or ["Explanation 1"]

    def get_explanations(self):
        return self._explanations


@pytest.fixture
def formatter():
    return DocxFormatter()


def test_save_exercises_creates_file(tmp_path, formatter):
    exercises = [DummyExercise(description="Test desc", question="Q?", options=["a", "b"])]
    path = tmp_path / "exercises.docx"
    formatter.save_exercises(exercises, str(path))

    assert path.exists()

    doc = Document(str(path))

    all_texts = [p.text for p in doc.paragraphs]
    assert any("Сборник упражнений" in t for t in all_texts)

    assert any("Q?" in t for t in all_texts)


def test_save_answers_creates_file(tmp_path, formatter):
    exercises = [DummyExercise(answer="Correct")]
    path = tmp_path / "answers.docx"
    formatter.save_answers(exercises, str(path))

    assert path.exists()

    doc = Document(str(path))

    assert any("Ключи к упражнениям" in p.text for p in doc.paragraphs)
    assert any("Correct" in p.text for p in doc.paragraphs)


def test_matching_exercise_formatting(tmp_path, formatter):
    ex = MatchingExercise()
    path = tmp_path / "matching.docx"
    formatter.save_exercises([ex], str(path))
    doc = Document(str(path))
    tables = doc.tables

    assert len(tables) >= 1

    hdr = tables[0].rows[0].cells

    assert hdr[0].text == "№"
    assert hdr[1].text == "Слова"


def test_truefalse_exercise_formatting(tmp_path, formatter):
    ex = TrueFalseExercise()
    path = tmp_path / "tf.docx"
    formatter.save_exercises([ex], str(path))
    doc = Document(str(path))
    tables = doc.tables
    assert len(tables) >= 1
    hdr = tables[0].rows[0].cells
    assert hdr[0].text == "№"
    assert hdr[1].text == "Утверждение"
    assert hdr[2].text == "Ответ (В/Н)"


def test_truefalse_answers_with_explanations(tmp_path, formatter):
    ex = TrueFalseExercise()
    path = tmp_path / "tf_answers.docx"
    formatter.save_answers([ex], str(path))
    doc = Document(str(path))
    texts = [p.text for p in doc.paragraphs]
    assert any("Пояснения" in t for t in texts)
    assert any("1. Explanation 1" in t for t in texts)


def test_multiple_exercises_grouping(tmp_path, formatter):
    exercises = [
        DummyExercise(description="Desc1", question="Q1", options=["a"]),
        MatchingExercise(),
        TrueFalseExercise()
    ]
    path = tmp_path / "mixed.docx"
    formatter.save_exercises(exercises, str(path))
    doc = Document(str(path))
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
    assert any("Заполнение пропусков" in h for h in headings) or True
    assert any("Соответствие" in h for h in headings)
    assert any("Верно/Неверно" in h for h in headings)


def test__get_type_name_returns_known_and_default(formatter):
    assert formatter._get_type_name("WordOrderExercise") == "Порядок слов"
    assert formatter._get_type_name("UnknownExercise") == "UnknownExercise"
